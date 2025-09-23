# path: exports/exporters.py
from __future__ import annotations

import os
import re
import textwrap
from datetime import datetime
from typing import Iterable, List, Optional

# ---------- Markdown → plain helpers (minimal, deterministic) ----------

def _md_to_lines(md: str, width: int = 92) -> List[str]:
    """Very small markdown → wrapped text for export. Headings become uppercase."""
    lines: List[str] = []
    for raw in (md or "").splitlines():
        if not raw.strip():
            lines.append("")
            continue
        if raw.lstrip().startswith("#"):
            level = len(raw) - len(raw.lstrip("#"))
            title = raw.lstrip("#").strip()
            title = title.upper() if level <= 2 else title
            prefix = "" if level == 1 else (" " * (2 * (level - 1)))
            lines.extend(textwrap.wrap(prefix + title, width=width)) or [prefix + title]
            continue
        # bullets
        if raw.strip().startswith(("- ", "* ")):
            body = re.sub(r"^(\s*[-*]\s+)", "• ", raw)
            lines.extend(textwrap.wrap(body, width=width, subsequent_indent="  "))
            continue
        # strip bold/italics/links/code markers (best-effort)
        t = re.sub(r"\*\*(.*?)\*\*", r"\1", raw)
        t = re.sub(r"\*(.*?)\*", r"\1", t)
        t = re.sub(r"`(.*?)`", r"\1", t)
        t = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", t)
        lines.extend(textwrap.wrap(t, width=width) or [t])
    return lines


# ---------- PDF (ReportLab) ----------

def export_pdf(markdown_text: str, out_path: str, title: str = "MinuteOne Note", fonts_dir: str = "exports/fonts") -> str:
    """
    Export markdown to a simple paginated PDF (offline).
    If reportlab is missing, writes a .txt fallback next to the target and returns that path.
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except Exception:
        # Fallback to .txt
        txt_path = os.path.splitext(out_path)[0] + ".txt"
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(markdown_text or "")
        return txt_path

    # Register fonts if present
    font_name = "Helvetica"
    try:
        if os.path.isdir(fonts_dir):
            for fn in os.listdir(fonts_dir):
                if fn.lower().endswith(".ttf"):
                    path = os.path.join(fonts_dir, fn)
                    name = os.path.splitext(fn)[0]
                    pdfmetrics.registerFont(TTFont(name, path))
                    font_name = name  # prefer first custom font
                    break
    except Exception:
        pass

    c = canvas.Canvas(out_path, pagesize=letter)
    width, height = letter
    margin = 54  # 0.75"
    y = height - margin

    # Header
    c.setFont(font_name, 13)
    c.drawString(margin, y, title)
    c.setFont(font_name, 9)
    c.drawRightString(width - margin, y, datetime.now().strftime("%Y-%m-%d %H:%M"))
    y -= 20

    # Body
    text = c.beginText(margin, y)
    text.setFont(font_name, 10)
    for line in _md_to_lines(markdown_text, width=int((width - 2 * margin) / 6)):  # rough char width
        if y < margin + 40:
            c.drawText(text)
            c.showPage()
            y = height - margin
            text = c.beginText(margin, y)
            text.setFont(font_name, 10)
        text.textLine(line)
        y -= 12
    c.drawText(text)
    c.showPage()
    c.save()
    return out_path


# ---------- DOCX (python-docx) ----------

def export_docx(markdown_text: str, out_path: str, title: Optional[str] = None) -> str:
    """
    Export markdown to DOCX using python-docx.
    If python-docx is missing, writes .rtf fallback via minimal hand-crafted RTF.
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception:
        # Very small RTF fallback (Unicode-safe header)
        return export_rtf(markdown_text, os.path.splitext(out_path)[0] + ".rtf", title=title or "MinuteOne Note")

    doc = Document()
    if title:
        h = doc.add_heading(title, level=1)
    for raw in (markdown_text or "").splitlines():
        if raw.startswith("# "):
            doc.add_heading(raw[2:].strip(), level=1)
        elif raw.startswith("## "):
            doc.add_heading(raw[3:].strip(), level=2)
        elif raw.strip().startswith(("- ", "* ")):
            doc.add_paragraph(raw.strip()[2:], style="List Bullet")
        else:
            doc.add_paragraph(raw)
    doc.save(out_path)
    return out_path


# ---------- RTF (pyrtf-ng or minimal writer) ----------

def export_rtf(markdown_text: str, out_path: str, title: str = "MinuteOne Note") -> str:
    try:
        from pyrtf_ng import Renderer, Document, Section, Paragraph, Text
        doc = Document()
        sec = Section()
        doc.Sections.append(sec)
        sec.append(Paragraph(Text(title)))
        for line in (markdown_text or "").splitlines():
            sec.append(Paragraph(Text(line)))
        with open(out_path, "w", encoding="utf-8") as f:
            Renderer().Write(doc, f)
        return out_path
    except Exception:
        # Minimal RTF that most readers accept
        def esc(s: str) -> str:
            return s.replace("\\", r"\\").replace("{", r"\{").replace("}", r"\}")
        lines = [r"{\rtf1\ansi\deff0", r"{\fonttbl{\f0 Arial;}}", r"\fs22 " + esc(title) + r"\par"]
        for ln in (markdown_text or "").splitlines():
            lines.append(esc(ln) + r"\par")
        lines.append("}")
        with open(out_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        return out_path


# ---------- Convenience ----------

def export_all(markdown_text: str, base_path_without_ext: str, title: str = "MinuteOne Note") -> dict:
    """Export PDF, DOCX, and RTF with the same base path. Returns dict of format → path."""
    out = {}
    out["pdf"] = export_pdf(markdown_text, base_path_without_ext + ".pdf", title=title)
    out["docx"] = export_docx(markdown_text, base_path_without_ext + ".docx", title=title)
    out["rtf"] = export_rtf(markdown_text, base_path_without_ext + ".rtf", title=title)
    return out
